import os
from typing import List

import docx
import pandas as pd


def read_all_text(document_path, sep="\n"):
    doc = docx.Document(document_path)
    final_text = []
    for para in doc.paragraphs:
        final_text.append(para.text)
    return sep.join(final_text)


def read_docx_table_to_list(document_path: str) -> List:
    doc = docx.Document(document_path)
    tables = doc.tables
    final_tables = []
    for table in tables:
        data = []
        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
        final_tables.append(data)
    return final_tables


def read_docx_table_to_df(document_path: str) -> List[pd.DataFrame]:
    tables = read_docx_table_to_list(document_path)
    tables_df = [pd.DataFrame.from_dict(table) for table in tables]
    return tables_df


def convert_docx_tables_to_xlsx(document_path: str):
    tables = read_docx_table_to_df(document_path)
    for count, table in enumerate(tables):
        converted_file_name = f"{os.path.splitext(document_path)[0]}-table-{count}.xlsx"
        print(f"Converting file {converted_file_name}...")
        table.to_excel(converted_file_name)


if __name__ == "__main__":
    # doc = docx.Document("Del_08_Generiranje_porocil/data/Uvod_v_Python.docx")
    # all_paras = doc.paragraphs

    # # # loop skozi vse paragraphe
    # # for para in all_paras:
    # #     print(para.text)
    # #     print("------")

    # sigle_para = doc.paragraphs[2]
    # print(sigle_para.text.lower())

    # tekst = read_all_text(
    #     "Del_08_Generiranje_porocil/data/Uvod_v_Python.docx", sep="**"
    # )
    # print(tekst)

    # Reading Runs
    # doc = docx.Document("Del_08_Generiranje_porocil/data/Uvod_v_Python.docx")
    # sigle_para = doc.paragraphs[2]

    # for run in sigle_para.runs:
    #     print(run.text)
    #     print("-----")
    # tables = read_docx_table_to_list(
    #     "Del_08_Generiranje_porocil/data/Uvod_v_Python.docx"
    # )
    # table1 = tables[0]

    # table2 = tables[1]
    # table2 = pd.DataFrame.from_dict(table2)
    # print(table2)

    # tables = read_docx_table_to_df("Del_08_Generiranje_porocil/data/Uvod_v_Python.docx")

    # print(tables[1])

    # read_docx_table_to_df()

    convert_docx_tables_to_xlsx("Del_08_Generiranje_porocil/data/Uvod_v_Python.docx")
