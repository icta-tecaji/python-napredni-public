import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from datetime import datetime


def write_new_doc(file_path: str):
    new_doc = docx.Document()
    new_doc.add_heading("Avtomatizirano poročilo", 0)

    timestamp = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    new_doc.add_paragraph(f"To poročilo je bilo zgenerirano ob {timestamp}")

    new_doc.add_heading("Del 1", 1)

    styles = new_doc.styles
    style = styles.add_style("tahoma_big", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Tahoma"
    style.font.size = Pt(25)

    p = new_doc.add_paragraph("Prvi del poročila vsebuje.")
    p.style = new_doc.styles["tahoma_big"]

    new_doc.save(file_path)


if __name__ == "__main__":
    write_new_doc("Del_08_Generiranje_porocil/data/novi_dokument.docx")
