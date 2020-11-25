import re

import docx
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill


class DocxTableReportParser:
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.workbook = Workbook()
        self.sheet = self.workbook.active
        self.tables = self._read_docx_table_all_elements_to_list(self.docx_path)
        # stili
        self.bold_font = Font(bold=True)
        self.center_aligned_text = Alignment(horizontal="center")
        self.light_gray_bg_color = PatternFill(start_color="D3D3D3", fill_type="solid")
        self.red_bg_color = PatternFill(start_color="FF0000", fill_type="solid")
        self.yellow_bg_color = PatternFill(start_color="FFFF00", fill_type="solid")
        self.cyan_bg_color = PatternFill(start_color="00FFFF", fill_type="solid")
        # dodamo prve dve vrstice
        self._add_first_line_template()
        self._add_second_line_template()

    def _append_rows(self, rows):
        for row in rows:
            self.sheet.append(row)

    def _get_class_id_version(self, table_id):
        data = self.tables[table_id][0][1].split(",")
        class_id = int(
            re.search(r"Class_id\s?=\s?(\d)", data[0], re.IGNORECASE).group(1)
        )
        version = int(re.search(r"version\s?=\s?(\d)", data[1], re.IGNORECASE).group(1))
        return class_id, version

    def _get_obis_code_daily(self, table_id):
        result = self.tables[table_id][2]
        result = [el.strip() for el in result]
        obis_code = (
            f"{result[1]}-{result[2]}:{result[3]}.{result[4]}.{result[5]}.{result[6]}"
        )
        return obis_code

    def _add_atributes_table(self, table_id):
        attributes = self.tables[table_id][3:]
        attributes_name = [
            [int(attr[0].split()[0].replace(".", "")), attr[0].split()[1]]
            for attr in attributes
        ]
        access_rights = []
        for ar in attributes:
            access_rights.append(
                [
                    el.replace("R/-", "Get_1")
                    .replace("R/W", "Get_1,Set_1")
                    .replace("-/- ", "")
                    for el in ar[2:]
                ]
            )
        final_list = []
        for an, ar in zip(attributes_name, access_rights):
            final_list.append([an[0], an[1], "", "", "", ar[0], ar[1], ar[3], ar[2]])
        self._append_rows(final_list)
        return len(attributes_name)

    def _read_docx_table_all_elements_to_list(self, document_path: str):
        doc = docx.Document(document_path)
        tables_doc = doc.tables
        final_tables = []
        for table in tables_doc:
            data = []
            for i, row in enumerate(table.rows):
                text = list((cell.text for cell in row.cells))
                data.append(text)
            final_tables.append(data)
        return final_tables

    def _add_first_line_template(self):
        """Add first line fix template."""
        rows = [
            [
                "",
                "Object/Attribute Name",
                "IC",
                "IC",
                "OBIS Object",
                "Access right",
                "Access right",
                "Access right",
                "Access right",
            ]
        ]

        self._append_rows(rows)

        # združimo celice
        self.sheet.merge_cells("C1:D1")
        self.sheet.merge_cells("F1:I1")
        # uredimo prvo vrstico
        for cell in self.sheet["1:1"]:
            cell.fill = self.light_gray_bg_color
            cell.alignment = self.center_aligned_text
            cell.font = self.bold_font

    def _add_second_line_template(self):
        """Add second line fix template"""
        rows = [
            [
                "#",
                "Object/Attribute Name",
                "Class ID",
                "Ver.",
                "OBIS Object Code / Default Value",
                "A.1",
                "A.2",
                "A.3",
                "A.4",
            ]
        ]
        self._append_rows(rows)
        for cell in self.sheet["2:2"]:
            cell.fill = self.light_gray_bg_color
            cell.alignment = self.center_aligned_text

    def add_energy_profile_daily_snapshot(self, id_class, id_attr):
        # dodamo rdečo vrstico
        rows = [["", "Energy Profile (Daily snapshot)"]]
        self._append_rows(rows)
        current_row = self.sheet._current_row
        for cell in self.sheet[f"{current_row}:{current_row}"]:
            cell.fill = self.red_bg_color
            cell.font = self.bold_font

        # dodamo rumeno vrstico
        calss_id, version = self._get_class_id_version(id_attr)
        obis_code = self._get_obis_code_daily(id_class)

        rows = [
            ["Attr", "Energy Profile (Daily snapshot)", calss_id, version, obis_code]
        ]
        self._append_rows(rows)
        current_row = self.sheet._current_row
        for cell in self.sheet[f"{current_row}:{current_row}"]:
            cell.fill = self.yellow_bg_color
            cell.font = self.bold_font

        # dodamo atribute
        table_len = self._add_atributes_table(id_attr)
        current_row = self.sheet._current_row
        for a in self.sheet[f"A{current_row-table_len+1}:B{current_row}"]:
            for cell in a:
                cell.fill = self.cyan_bg_color

        # dodamo prazno vrstico
        self.sheet._current_row += 1

    def add_total_energy_registers(self, id_class, id_attr):
        # dodamo rdečo vrstico
        rows = [["", "Total Energy Registers"]]
        self._append_rows(rows)
        current_row = self.sheet._current_row
        for cell in self.sheet[f"{current_row}:{current_row}"]:
            cell.fill = self.red_bg_color
            cell.font = self.bold_font

        # dodamo rumeno vrstico
        calss_id, version = self._get_class_id_version(id_attr)
        obis_code = self.tables[0][1][-1].strip()

        rows = [["Attr", "Active energy import (+A)", calss_id, version, obis_code]]
        self._append_rows(rows)
        current_row = self.sheet._current_row
        for cell in self.sheet[f"{current_row}:{current_row}"]:
            cell.fill = self.yellow_bg_color
            cell.font = self.bold_font

        # dodamo atribute
        table_len = self._add_atributes_table(id_attr)
        current_row = self.sheet._current_row
        for a in self.sheet[f"A{current_row-table_len+1}:B{current_row}"]:
            for cell in a:
                cell.fill = self.cyan_bg_color

        # dodamo prazno vrstico
        self.sheet._current_row += 1

        # dodamo rumeno vrstico
        calss_id, version = self._get_class_id_version(id_attr)
        obis_code = self.tables[0][2][-1].strip()

        rows = [["Attr", "Active energy import (-A)", calss_id, version, obis_code]]
        self._append_rows(rows)
        current_row = self.sheet._current_row
        for cell in self.sheet[f"{current_row}:{current_row}"]:
            cell.fill = self.yellow_bg_color
            cell.font = self.bold_font

        # dodamo atribute
        table_len = self._add_atributes_table(id_attr)
        current_row = self.sheet._current_row
        for a in self.sheet[f"A{current_row-table_len+1}:B{current_row}"]:
            for cell in a:
                cell.fill = self.cyan_bg_color

        # dodamo prazno vrstico
        self.sheet._current_row += 1

    def save_table(self, table_path):
        self.workbook.save(filename=table_path)


if __name__ == "__main__":
    my_parser = DocxTableReportParser(
        "Dodatni_primeri/data/PRIMER_kompleksna_tabela.docx"
    )

    my_parser.add_total_energy_registers(0, 1)
    my_parser.add_energy_profile_daily_snapshot(2, 3)
    my_parser.save_table("data/complex_report.xlsx")
